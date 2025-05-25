from docx import Document
import os

os.makedirs("sample_inputs", exist_ok=True)

doc = Document()
doc.add_paragraph("We had a great meeting today.")
doc.add_paragraph("- John should finalize the Q2 sales report by June 10.")
doc.add_paragraph("- Lisa to prepare the client presentation next Monday.")
doc.add_paragraph("- Schedule the internal review session before Friday.")
doc.save("sample_inputs/sample_notes.docx")
print("✅ sample_notes.docx created")
